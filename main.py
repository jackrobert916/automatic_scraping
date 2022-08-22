import docx
from docx.shared import Inches, Mm
import os
from PIL import Image
from PIL import Image, ImageDraw, ImageFont
from docx2pdf import convert
from wand.image import Image as Im

image_dir = os.listdir(os.getcwd()+'\\Images\\Images')
print(len(image_dir))
doc = docx.Document()
section = doc.sections[0]
section.page_height = Mm(1000)
section.page_width = Mm(580)
section.left_margin = Mm(25.4)
section.right_margin = Mm(25.4)
section.top_margin = Mm(25.4)
section.bottom_margin = Mm(25.4)
section.header_distance = Mm(12.7)
section.footer_distance = Mm(12.7)
p = doc.add_paragraph()
x = 0
for i in range(0, len(image_dir)):
    size = (130, 160)
    temp_img = Image.open(os.getcwd()+'\\Images\\Images\\'+image_dir[i])
    temp_img = temp_img.resize(size)
    # temp_img.thumbnail(size, Image.ANTIALIAS)
    
    # temp_img.show()
    background = Image.new('RGBA', (500, 220), (255, 255, 255, 0))
    for k in range(0, 3):
        background.paste(temp_img, (0,0))
        background.paste(temp_img, (150,0))
        background.paste(temp_img, (300,0))
    font = ImageFont.truetype(r'arial.ttf', 25) 
    d1 = ImageDraw.Draw(background)
    d1.text((5, 160), image_dir[i][:-4], fill =(0, 0, 0), font = font)
    background.save("temp.png")
    with Im(filename ="temp.png") as img:
    
        # generating sharp image using sharpen() function.
        img.sharpen(radius = 16, sigma = 8)
        img.save(filename ="temp1.png")
    r = p.add_run()
    r.add_picture("temp1.png")
    # if x == 2:
    #     p = doc.add_paragraph()
    #     x = 0
    # else:
    #     x+=1
    #     continue
        
doc.save('demo1.docx')
convert("demo1.docx")
# convert("demodocx", "result.pdf")
